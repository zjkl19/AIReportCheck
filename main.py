# -*- coding: utf-8 -*-
"""
Created on 2019/05/25

@author: lindinan
"""

#TODO:
#doc格式用win32com转为docx

#仪器有旧编号的要使用旧编号
#方案要有人员
#ReportError等添加修改建议


import docx
import AIReportCheck


docName='hyl.docx'

#import os
#import win32com
#from win32com.client import Dispatch
##导入库
#wordApp = win32com.client.Dispatch('Word.Application')
##调用word程序
#wordApp.Visible = 0
#wordApp.DisplayAlerts = 0
##不在前台显示文档及错误，在实际使用阶段可以全部关闭，提高运行速度，但是在调试时打开还是用处挺大的，可以对操作是否实现自己的需求进行直观的判断，比如说我们选中的内容是否已经高亮等等。
#doc = wordApp.Documents.Open(u''+os.getcwd()+'\\'+docName)
#try:
#    document = wordApp.Documents.Open(u''+os.getcwd()+'\\'+docName)
#except Exception as e:
#    print(e.message)

try:
    document = docx.Document(docName)  #打开文件
except Exception as e:
    print(e.message)

r=AIReportCheck.AIReportCheck(document)
r.CheckReport()

for e in r.ErrorList:
    print(e.Name)

import tkinter as tk
from tkinter import ttk

win = tk.Tk()
win.title("Python GUI")    # 添加标题

ttk.Label(win, text="Chooes a number").grid(column=1, row=0)    # 添加一个标签，并将其列设置为1，行设置为0
ttk.Label(win, text="Enter a name:").grid(column=0, row=0)      # 设置其在界面中出现的位置  column代表列   row 代表行

# button被点击之后会被执行
def clickMe():   # 当acction被点击时,该函数则生效
  action.configure(text='Hello ' + name.get() + ' ' + numberChosen.get())     # 设置button显示的内容
  print('check3 is %s %s' % (type(chvarEn.get()), chvarEn.get()))

# 按钮
action = ttk.Button(win, text="Click Me!", command=clickMe)     # 创建一个按钮, text：显示按钮上面显示的文字, command：当这个按钮被点击之后会调用command函数
action.grid(column=2, row=1)    # 设置其在界面中出现的位置  column代表列   row 代表行

# 文本框
name = tk.StringVar()     # StringVar是Tk库内部定义的字符串变量类型，在这里用于管理部件上面的字符；不过一般用在按钮button上。改变StringVar，按钮上的文字也随之改变。
nameEntered = ttk.Entry(win, width=12, textvariable=name)   # 创建一个文本框，定义长度为12个字符长度，并且将文本框中的内容绑定到上一句定义的name变量上，方便clickMe调用
nameEntered.grid(column=0, row=1)       # 设置其在界面中出现的位置  column代表列   row 代表行
nameEntered.focus()     # 当程序运行时,光标默认会出现在该文本框中

# 创建一个下拉列表
number = tk.StringVar()
numberChosen = ttk.Combobox(win, width=12, textvariable=number, state='readonly')
numberChosen['values'] = (1, 2, 4, 42, 100)     # 设置下拉列表的值
numberChosen.grid(column=1, row=1)      # 设置其在界面中出现的位置  column代表列   row 代表行
numberChosen.current(0)    # 设置下拉列表默认显示的值，0为 numberChosen['values'] 的下标值

# 复选框
chVarDis = tk.IntVar()   # 用来获取复选框是否被勾选，通过chVarDis.get()来获取其的状态,其状态值为int类型 勾选为1  未勾选为0
check1 = tk.Checkbutton(win, text="Disabled", variable=chVarDis, state='disabled')    # text为该复选框后面显示的名称, variable将该复选框的状态赋值给一个变量，当state='disabled'时，该复选框为灰色，不能点的状态
check1.select()     # 该复选框是否勾选,select为勾选, deselect为不勾选
check1.grid(column=0, row=4, sticky=tk.W)       # sticky=tk.W  当该列中其他行或该行中的其他列的某一个功能拉长这列的宽度或高度时，设定该值可以保证本行保持左对齐，N：北/上对齐  S：南/下对齐  W：西/左对齐  E：东/右对齐

chvarUn = tk.IntVar()
check2 = tk.Checkbutton(win, text="UnChecked", variable=chvarUn)
check2.deselect()
check2.grid(column=1, row=4, sticky=tk.W)

chvarEn = tk.IntVar()
check3 = tk.Checkbutton(win, text="Enabled", variable=chvarEn)
check3.select()
check3.grid(column=2, row=4, sticky=tk.W)

win.mainloop()      # 当调用mainloop()时,窗口才会显示出来

#try:
#    doc.Close()
#    wordApp.Quit()
#except Exception as e:
#    print(e.message)

#tables = document.tables #获取文件中的表格集
#table = tables[0]#获取文件中的第一个表格
#result=""
#for i in range(0,len(table.rows)):#从表格第二行开始循环读取表格数据
#    result = table.cell(i,j).text# + "" +table.cell(i,1).text+table.cell(i,2).text + table.cell(i,3).text
#    #cell(i,0)表示第(i+1)行第1列数据，以此类推
#    print(result)

#totalText=result+mainBody
#import re
#ret = re.findall(r'(Km/h)',totalText)
#if len(ret) > 0:
#   print("速度单位疑似存在错误")

#for i in range(0,len(document.paragraphs)):   
    #print(document.paragraphs[i].text)



